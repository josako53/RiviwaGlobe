"""generate_pricing_doc.py — Generate Riviwa subscription pricing DOCX design document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(0.9)
section.right_margin  = Inches(0.9)
section.top_margin    = Inches(0.8)
section.bottom_margin = Inches(0.8)

# ── Brand colors ──────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1a, 0x1a, 0x2e)   # Riviwa primary
BLUE        = RGBColor(0x1d, 0x4e, 0xd8)   # verified blue
GREEN       = RGBColor(0x16, 0xa3, 0x4a)   # active green
PURPLE      = RGBColor(0x7c, 0x3a, 0xed)   # professional
ORANGE      = RGBColor(0xea, 0x58, 0x0c)   # business
GOLD        = RGBColor(0xd9, 0x77, 0x06)   # enterprise
LIGHT_GREY  = RGBColor(0xf3, 0xf4, 0xf6)
MID_GREY    = RGBColor(0x6b, 0x72, 0x80)
DARK_GREY   = RGBColor(0x37, 0x41, 0x51)
WHITE       = RGBColor(0xff, 0xff, 0xff)
TABLE_HEAD  = RGBColor(0x1a, 0x1a, 0x2e)
ROW_ALT     = RGBColor(0xf8, 0xf9, 0xff)
GREEN_LIGHT = RGBColor(0xdc, 0xfc, 0xe7)
BLUE_LIGHT  = RGBColor(0xdb, 0xea, 0xfe)
RED_LIGHT   = RGBColor(0xfe, 0xe2, 0xe2)

# ── Helper functions ──────────────────────────────────────────────────────────

def rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(color))
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), rgb_hex(color))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def cell_para(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    return para

def add_heading(doc, text, level=1, color=NAVY, size=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    sizes = {1: 20, 2: 16, 3: 13, 4: 11}
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.color.rgb = color
    return p

def add_para(doc, text, size=10, color=DARK_GREY, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.bold   = bold
    run.italic = italic
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Top accent bar (via table)
bar = doc.add_table(rows=1, cols=1)
bar.alignment = WD_TABLE_ALIGNMENT.CENTER
bar_cell = bar.rows[0].cells[0]
set_cell_bg(bar_cell, NAVY)
bar_cell.height = Cm(1.2)
cell_para(bar_cell, 'RIVIWA PLATFORM', bold=True, size=11, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_heading(doc, 'Subscription Plans &\nPricing Design Guide', level=1, size=28, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=6)
add_para(doc, 'Platform Scope · Organisation Scope · Checkout Flow', size=12, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=20)

# Subtitle band
band = doc.add_table(rows=1, cols=1)
band.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = band.rows[0].cells[0]
set_cell_bg(bc, BLUE_LIGHT)
cell_para(bc, 'Riviwa Grievance & Feedback Management Platform  ·  Version 2.5  ·  May 2026', bold=False, size=10, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

# Three plan summary cards (single-row table)
card_tbl = doc.add_table(rows=2, cols=3)
card_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

plans_summary = [
    ('Starter',      '$15 / mo',  BLUE,   'For small NGOs & CBOs'),
    ('Professional', '$49 / mo',  PURPLE, 'For hospitals & banks'),
    ('Business',     '$149 / mo', ORANGE, 'Full platform access'),
]

for i, (name, price, color, desc) in enumerate(plans_summary):
    c = card_tbl.rows[0].cells[i]
    set_cell_bg(c, color)
    cell_para(c, name,  bold=True, size=13, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (name, price, color, desc) in enumerate(plans_summary):
    c = card_tbl.rows[1].cells[i]
    set_cell_bg(c, LIGHT_GREY)
    cell_para(c, f'{price}\n{desc}', bold=False, size=9, color=DARK_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para(doc, 'Plus Enterprise (custom pricing) · 14-day free trial on all plans · No credit card required',
         size=9, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Contents', level=2, color=NAVY, space_before=10, space_after=10)
add_divider(doc)

toc_items = [
    ('1', 'Platform Admin — Plan Overview',               'Seeded plans, pricing, limits'),
    ('2', 'Platform Admin — Feature Matrix',              'All 46 features across 4 plans'),
    ('3', 'Platform Admin — Promo Codes & Sales',         'Active codes and campaigns'),
    ('4', 'Org User — Pricing Page Design',               'Shopify-style layout with endpoints'),
    ('5', 'Org User — Plan Comparison Table',             'Feature-by-feature comparison'),
    ('6', 'Org User — Checkout Step-by-Step',             'From plan selection to payment'),
    ('7', 'Org User — Payment Provider Flows',            'M-Pesa, PayPal, bank transfer'),
    ('8', 'Org User — Post-Payment & Verification',       'Badge, features unlocked, KYC'),
    ('9', 'Endpoint Reference per Screen',                'Full API map for each page'),
]

toc_tbl = doc.add_table(rows=len(toc_items), cols=3)
toc_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (num, title, desc) in enumerate(toc_items):
    cells = toc_tbl.rows[i].cells
    set_cell_bg(cells[0], NAVY if i % 2 == 0 else DARK_GREY)
    cell_para(cells[0], num, bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(cells[1], title, bold=True, size=10, color=NAVY)
    cell_para(cells[2], desc, bold=False, size=9, color=MID_GREY, italic=True)
    cells[0].width = Cm(1.0)
    cells[1].width = Cm(8.0)
    cells[2].width = Cm(7.5)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — PLATFORM ADMIN: PLAN OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '1.  Platform Admin — Plan Overview', level=2, color=NAVY)
add_divider(doc)
add_para(doc, 'The following plans are seeded at platform startup. Platform admins access these via the Billing Dashboard. Organisations see only public plans on the pricing page.', size=9, color=MID_GREY, space_before=2, space_after=8)

# Plan overview table
headers = ['', 'Starter', 'Professional', 'Business', 'Enterprise']
plan_rows = [
    ('Monthly price',            '$15.00',    '$49.00',    '$149.00',   'Custom'),
    ('Annual price / mo',        '$12.00',    '$39.00',    '$119.00',   'Custom'),
    ('Annual total',             '$144 / yr', '$468 / yr', '$1,428 / yr','Negotiated'),
    ('Annual saving vs monthly', '20%',       '20%',       '20%',       '—'),
    ('Free trial',               '14 days',   '14 days',   '14 days',   '30 days'),
    ('Team members',             '5',         '25',        '100',       'Unlimited'),
    ('Projects',                 '3',         '15',        'Unlimited', 'Unlimited'),
    ('Submissions / month',      '500',       '5,000',     'Unlimited', 'Unlimited'),
    ('SMS / month',              '200',       '2,000',     '10,000',    'Unlimited'),
    ('API calls / month',        '2,000',     '10,000',    '100,000',   'Unlimited'),
    ('Storage',                  '5 GB',      '25 GB',     '100 GB',    'Unlimited'),
    ('QR codes / month',         '50',        '500',       'Unlimited', 'Unlimited'),
    ('Staff profiles',           '—',         '100',       'Unlimited', 'Unlimited'),
    ('Uptime SLA',               '99.5%',     '99.9%',     '99.95%',    '99.99%'),
    ('Support',                  'Email',     'Priority email + chat', 'Dedicated CSM', 'Enterprise CSM'),
    ('Custom domain / SSO',      '—',         '—',         '—',         '✓'),
    ('White-label',              '—',         '—',         '—',         '✓'),
    ('Custom SLA contract',      '—',         '—',         '—',         '✓'),
    ('Is public (pricing page)', '✓ Yes',     '✓ Yes',     '✓ Yes',     '✓ Yes'),
    ('Sort order',               '1',         '2',         '3',         '4'),
]

col_colors = [NAVY, BLUE, PURPLE, ORANGE, GOLD]

pt = doc.add_table(rows=1 + len(plan_rows), cols=5)
pt.style = 'Table Grid'
pt.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for j, h in enumerate(headers):
    c = pt.rows[0].cells[j]
    set_cell_bg(c, col_colors[j])
    cell_para(c, h, bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

# Data rows
for i, (label, *vals) in enumerate(plan_rows):
    row_bg = LIGHT_GREY if i % 2 == 0 else WHITE
    cells = pt.rows[i + 1].cells
    set_cell_bg(cells[0], DARK_GREY if i % 2 == 0 else RGBColor(0x4b, 0x55, 0x63))
    cell_para(cells[0], label, bold=True, size=9, color=WHITE)
    for j, val in enumerate(vals, 1):
        set_cell_bg(cells[j], row_bg)
        txt_color = GREEN if val.startswith('✓') else (MID_GREY if val == '—' else DARK_GREY)
        cell_para(cells[j], val, bold=False, size=9, color=txt_color, align=WD_ALIGN_PARAGRAPH.CENTER)

# Column widths
widths = [4.0, 2.8, 2.8, 2.8, 2.8]
for j, w in enumerate(widths):
    set_col_width(pt, j, w)

add_para(doc, '* -1 in API means unlimited. All prices in USD.', size=8, color=MID_GREY, italic=True, space_before=4)

# Endpoints box
add_para(doc, 'Admin endpoints for this section:', bold=True, size=9, color=NAVY, space_before=10)
ep_tbl = doc.add_table(rows=4, cols=2)
eps = [
    ('GET  /api/v1/plans/admin/plans',                  'List all plans including inactive'),
    ('PATCH /api/v1/plans/admin/plans/{id}/pricing',    'Update plan pricing'),
    ('PATCH /api/v1/plans/admin/plans/{id}/limits',     'Update usage limits'),
    ('GET  /api/v1/billing/metrics',                    'MRR, ARR, plan distribution, churn'),
]
for i, (ep, desc) in enumerate(eps):
    set_cell_bg(ep_tbl.rows[i].cells[0], NAVY)
    cell_para(ep_tbl.rows[i].cells[0], ep, bold=False, size=8, color=WHITE)
    set_cell_bg(ep_tbl.rows[i].cells[1], ROW_ALT)
    cell_para(ep_tbl.rows[i].cells[1], desc, bold=False, size=8, color=DARK_GREY)
    ep_tbl.rows[i].cells[0].width = Cm(8.0)
    ep_tbl.rows[i].cells[1].width = Cm(8.5)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — PLATFORM ADMIN: FEATURE MATRIX
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '2.  Platform Admin — Feature Matrix', level=2, color=NAVY)
add_divider(doc)
add_para(doc, 'All 46 features grouped by service category. ✓ = included in plan. — = not included.', size=9, color=MID_GREY, space_before=2, space_after=8)

TICK = '✓'
DASH = '—'

feature_groups = [
    ('Feedback Channels',       [
        ('SMS Channel',                    TICK, TICK, TICK, TICK),
        ('WhatsApp Channel',               DASH, TICK, TICK, TICK),
        ('Phone Call AI (IVR)',            DASH, DASH, TICK, TICK),
    ]),
    ('AI & Intelligence',       [
        ('AI Conversation (Web/Mobile)',   DASH, TICK, TICK, TICK),
        ('AI-Powered Insights',           DASH, DASH, TICK, TICK),
        ('Voice Transcription',           DASH, TICK, TICK, TICK),
        ('ML Escalation Predictor',       DASH, DASH, TICK, TICK),
        ('Real-Time Spark Streaming',     DASH, DASH, TICK, TICK),
        ('Smart Recommendations',         TICK, TICK, TICK, TICK),
        ('AI Counterfeit Analysis',       DASH, TICK, TICK, TICK),
    ]),
    ('Notifications',           [
        ('Push Notifications',            TICK, TICK, TICK, TICK),
        ('WhatsApp Notifications',        DASH, TICK, TICK, TICK),
    ]),
    ('Analytics',               [
        ('Advanced Analytics Dashboard',  DASH, TICK, TICK, TICK),
        ('Custom Report Builder',         DASH, TICK, TICK, TICK),
    ]),
    ('Feedback Management',     [
        ('Employee Feedback (360°)',      TICK, TICK, TICK, TICK),
        ('PAP Registry',                  TICK, TICK, TICK, TICK),
        ('Committee Management',          DASH, TICK, TICK, TICK),
        ('Bulk Feedback Import',          DASH, TICK, TICK, TICK),
    ]),
    ('QR & Verification',       [
        ('QR Code Generation',            TICK, TICK, TICK, TICK),
        ('Product Verification',          TICK, TICK, TICK, TICK),
        ('Field Agent Management',        DASH, TICK, TICK, TICK),
    ]),
    ('Staff',                   [
        ('Staff Identity Verification',   TICK, TICK, TICK, TICK),
        ('Bulk Staff CSV Import',         DASH, TICK, TICK, TICK),
        ('Staff Analytics',               DASH, TICK, TICK, TICK),
    ]),
    ('Queue Management',        [
        ('Multi-Step Queue Management',   DASH, TICK, TICK, TICK),
    ]),
    ('Stakeholder',             [
        ('Stakeholder Engagement (SEP)',  DASH, TICK, TICK, TICK),
    ]),
    ('Translation',             [
        ('Auto-Translation (63 lang.)',   TICK, TICK, TICK, TICK),
        ('Advanced Translation',          DASH, TICK, TICK, TICK),
    ]),
    ('Product Catalog',         [
        ('Product Catalog',               TICK, TICK, TICK, TICK),
        ('Product Variations',            DASH, TICK, TICK, TICK),
        ('RSIN (Item Numbering)',          DASH, TICK, TICK, TICK),
    ]),
    ('Integration',             [
        ('REST API Access',               TICK, TICK, TICK, TICK),
        ('Webhook Engine',                DASH, TICK, TICK, TICK),
        ('OAuth2 PKCE',                   TICK, TICK, TICK, TICK),
        ('JavaScript Widget Embed',       TICK, TICK, TICK, TICK),
        ('Audit Logs',                    DASH, TICK, TICK, TICK),
    ]),
    ('Payments',                [
        ('Mobile Money (TZ)',             DASH, TICK, TICK, TICK),
        ('PayPal',                        DASH, DASH, TICK, TICK),
        ('Payment Processing (full)',     DASH, TICK, TICK, TICK),
    ]),
    ('Authentication',          [
        ('Social Login',                  TICK, TICK, TICK, TICK),
        ('ID Verification',               DASH, TICK, TICK, TICK),
        ('Fraud Detection',               TICK, TICK, TICK, TICK),
        ('Multi-Org Switching',           TICK, TICK, TICK, TICK),
        ('Two-Factor Authentication',     TICK, TICK, TICK, TICK),
        ('Single Sign-On (SSO)',          DASH, DASH, DASH, TICK),
    ]),
    ('Platform',                [
        ('White-Label',                   DASH, DASH, DASH, TICK),
        ('Dedicated Support',             DASH, TICK, TICK, TICK),
        ('Custom SLA (99.99%)',           DASH, DASH, DASH, TICK),
        ('Geo Proximity',                 DASH, TICK, TICK, TICK),
        ('Waiting Queue',                 DASH, TICK, TICK, TICK),
    ]),
]

fm_headers = ['Feature', 'Starter', 'Professional', 'Business', 'Enterprise']
fm_col_colors = [NAVY, BLUE, PURPLE, ORANGE, GOLD]

fm = doc.add_table(rows=1, cols=5)
fm.style = 'Table Grid'
fm.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, h in enumerate(fm_headers):
    c = fm.rows[0].cells[j]
    set_cell_bg(c, fm_col_colors[j])
    cell_para(c, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

row_idx = 1
for group_name, features in feature_groups:
    # Group header row
    r = fm.add_row()
    for j in range(5):
        set_cell_bg(r.cells[j], DARK_GREY)
    cell_para(r.cells[0], group_name, bold=True, size=9, color=WHITE)
    for j in range(1, 5):
        cell_para(r.cells[j], '', size=9, color=WHITE)
    row_idx += 1

    for feat_name, s, p, b, e in features:
        r = fm.add_row()
        row_bg = LIGHT_GREY if row_idx % 2 == 0 else WHITE
        set_cell_bg(r.cells[0], row_bg)
        cell_para(r.cells[0], feat_name, bold=False, size=8.5, color=DARK_GREY)
        for j, val in enumerate([s, p, b, e], 1):
            set_cell_bg(r.cells[j], GREEN_LIGHT if val == TICK else row_bg)
            cell_para(r.cells[j], val, bold=(val == TICK), size=9,
                      color=GREEN if val == TICK else MID_GREY,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        row_idx += 1

for j, w in enumerate([5.5, 2.3, 2.3, 2.3, 2.3]):
    set_col_width(fm, j, w)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — PLATFORM ADMIN: PROMO CODES & SALES
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '3.  Platform Admin — Promo Codes & Active Sales', level=2, color=NAVY)
add_divider(doc)

add_para(doc, 'SEEDED PROMO CODES', bold=True, size=10, color=NAVY, space_before=8, space_after=4)

promo_headers = ['Code', 'Discount', 'Duration', 'Eligible Plans', 'Limit', 'Expires']
promos = [
    ('LAUNCH2026', '30% off',     '3 months',  'All plans',              'Unlimited', '31 Dec 2026'),
    ('ANNUAL20',   '20% off',     'Once',      'All plans (annual)',     'Unlimited', 'No expiry'),
    ('NGO50',      '50% off',     'Forever',   'Starter only',          '100 uses',  'No expiry'),
    ('PARTNER25',  '25% off',     '6 months',  'Professional / Business','500 uses',  '31 Dec 2026'),
    ('WELCOME1',   '1 free month','Once',      'All plans',              'Unlimited', '30 Sep 2026'),
    ('GOV30',      '30% off',     'Forever',   'Pro / Business / Ent.',  '200 uses',  'No expiry'),
    ('UPGRADE30',  '30% off',     'Once',      'Business only',          'Unlimited', '31 Dec 2026'),
    ('STUDENT15',  '15% off',     '12 months', 'Professional only',      '300 uses',  '30 Jun 2027'),
    ('EARLYBIRD40','40% off',     '2 months',  'All plans',              '50 uses',   '1 Aug 2026'),
    ('DONOR10',    '$10 off',     'Forever',   'Pro / Business',         'Unlimited', 'No expiry'),
]

pt2 = doc.add_table(rows=1 + len(promos), cols=6)
pt2.style = 'Table Grid'
for j, h in enumerate(promo_headers):
    c = pt2.rows[0].cells[j]
    set_cell_bg(c, NAVY)
    cell_para(c, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row_data in enumerate(promos):
    row_bg = GREEN_LIGHT if i % 2 == 0 else WHITE
    for j, val in enumerate(row_data):
        c = pt2.rows[i+1].cells[j]
        set_cell_bg(c, row_bg)
        is_code = j == 0
        cell_para(c, val, bold=is_code, size=8.5,
                  color=NAVY if is_code else DARK_GREY,
                  align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT)

for j, w in enumerate([2.5, 2.0, 2.0, 3.5, 2.0, 2.5]):
    set_col_width(pt2, j, w)

doc.add_paragraph()
add_para(doc, 'ACTIVE SALES CAMPAIGNS', bold=True, size=10, color=NAVY, space_before=10, space_after=4)

sales_headers = ['Campaign', 'Discount', 'Auto-apply', 'Start', 'End', 'Applies to']
sales = [
    ('Launch Week Sale',     '40% off', 'Yes', '17 May 2026', '31 May 2026', 'All plans — new subscribers'),
    ('Mid-Year Sale 2026',   '25% off', 'Yes', '1 Jun 2026',  '30 Jun 2026', 'All plans'),
    ('Annual Plan Bonus',    '15% off', 'Yes', '1 Jul 2026',  '31 Aug 2026', 'Annual billing only'),
    ('Black Friday 2026',    '50% off', 'Yes', '27 Nov 2026', '27 Nov 2026', 'All plans'),
    ('Cyber Monday 2026',    '35% off', 'Yes', '30 Nov 2026', '30 Nov 2026', 'Pro / Business'),
    ('New Year 2027',        '20% off', 'Yes', '1 Jan 2027',  '14 Jan 2027', 'All plans'),
]

st = doc.add_table(rows=1 + len(sales), cols=6)
st.style = 'Table Grid'
for j, h in enumerate(sales_headers):
    c = st.rows[0].cells[j]
    set_cell_bg(c, DARK_GREY)
    cell_para(c, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row_data in enumerate(sales):
    row_bg = BLUE_LIGHT if i % 2 == 0 else WHITE
    for j, val in enumerate(row_data):
        c = st.rows[i+1].cells[j]
        set_cell_bg(c, row_bg)
        cell_para(c, val, bold=(j==0), size=8.5,
                  color=NAVY if j==0 else (GREEN if val=='Yes' else DARK_GREY),
                  align=WD_ALIGN_PARAGRAPH.LEFT if j in (0,5) else WD_ALIGN_PARAGRAPH.CENTER)
for j, w in enumerate([3.5, 2.0, 2.0, 2.5, 2.5, 4.0]):
    set_col_width(st, j, w)

add_para(doc, 'Admin endpoints: POST /api/v1/sales/admin  |  PATCH /api/v1/sales/admin/{id}  |  GET /api/v1/promotions/admin',
         size=8, color=MID_GREY, italic=True, space_before=6)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — ORG USER: PRICING PAGE DESIGN (Shopify-style)
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '4.  Org User — Pricing Page Design', level=2, color=NAVY)
add_divider(doc)
add_para(doc, 'Shopify-style pricing page layout. Toggle between monthly and annual billing. Cards highlight the recommended plan. Enterprise shows "Contact us" instead of a price.', size=9, color=MID_GREY, space_before=2, space_after=10)

# Page wireframe mockup using tables

# Header bar of pricing page
hdr = doc.add_table(rows=1, cols=1)
hdr_cell = hdr.rows[0].cells[0]
set_cell_bg(hdr_cell, NAVY)
cell_para(hdr_cell, 'riviwa.com / pricing', bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# Sale banner
banner = doc.add_table(rows=1, cols=1)
bc2 = banner.rows[0].cells[0]
set_cell_bg(bc2, GREEN)
cell_para(bc2, '🎉  Launch Week — 40% OFF for new subscribers. Ends 31 May 2026!', bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# Page title
add_heading(doc, 'Simple, transparent pricing', level=1, color=NAVY, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
add_para(doc, 'Start free for 14 days · No credit card required · Cancel anytime', size=9, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)

# Billing toggle
toggle = doc.add_table(rows=1, cols=3)
toggle.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(toggle.rows[0].cells[0], LIGHT_GREY)
cell_para(toggle.rows[0].cells[0], '', size=9, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell_bg(toggle.rows[0].cells[1], NAVY)
cell_para(toggle.rows[0].cells[1], '  Monthly    ●    Annual (save 20%)  ', bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(toggle.rows[0].cells[2], LIGHT_GREY)
cell_para(toggle.rows[0].cells[2], '', size=9, color=MID_GREY)
toggle.rows[0].cells[0].width = Cm(3.0)
toggle.rows[0].cells[1].width = Cm(8.0)
toggle.rows[0].cells[2].width = Cm(3.0)

doc.add_paragraph()

# Plan cards (2-row table: card header + card body)
plans_card_data = [
    {
        'name': 'Starter',
        'color': BLUE,
        'price_mo': '$15',
        'price_yr': '$12',
        'tagline': 'For small NGOs & CBOs',
        'popular': False,
        'cta': 'Start free trial',
        'highlights': [
            '5 team members',
            '3 projects',
            '500 submissions / mo',
            '200 SMS / mo',
            '2,000 API calls',
            '5 GB storage',
            '50 QR codes / mo',
            '14-day free trial',
        ],
    },
    {
        'name': 'Professional',
        'color': PURPLE,
        'price_mo': '$49',
        'price_yr': '$39',
        'tagline': 'For hospitals & banks',
        'popular': True,
        'cta': 'Start free trial',
        'highlights': [
            '25 team members',
            '15 projects',
            '5,000 submissions / mo',
            '2,000 SMS + WhatsApp',
            '10,000 API calls',
            '25 GB storage',
            '500 QR codes / mo',
            'AI Conversation',
            'Advanced Analytics',
            'Webhooks + Audit Logs',
        ],
    },
    {
        'name': 'Business',
        'color': ORANGE,
        'price_mo': '$149',
        'price_yr': '$119',
        'tagline': 'Full platform · All 15 services',
        'popular': False,
        'cta': 'Start free trial',
        'highlights': [
            '100 team members',
            'Unlimited projects',
            'Unlimited submissions',
            '10,000 SMS + WhatsApp',
            '100,000 API calls',
            '100 GB storage',
            'Unlimited QR codes',
            'Phone Call AI (IVR)',
            'Spark Streaming + ML',
            'Mobile Money + PayPal',
        ],
    },
    {
        'name': 'Enterprise',
        'color': GOLD,
        'price_mo': 'Custom',
        'price_yr': 'Custom',
        'tagline': 'For governments & UN agencies',
        'popular': False,
        'cta': 'Contact us',
        'highlights': [
            'Unlimited everything',
            'SSO + SCIM',
            'White-label',
            'Custom SLA 99.99%',
            'On-premise option',
            'Dedicated CSM',
            'Custom AI fine-tuning',
            'Volume SMS pricing',
        ],
    },
]

# Card header row
card_hdr = doc.add_table(rows=1, cols=4)
card_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, plan in enumerate(plans_card_data):
    c = card_hdr.rows[0].cells[i]
    set_cell_bg(c, plan['color'])
    text = f"{'★ MOST POPULAR\n' if plan['popular'] else ''}{plan['name']}"
    cell_para(c, text, bold=True, size=11, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    c.width = Cm(3.8)

# Card pricing row
card_price = doc.add_table(rows=1, cols=4)
card_price.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, plan in enumerate(plans_card_data):
    c = card_price.rows[0].cells[i]
    bg = BLUE_LIGHT if plan['color'] == BLUE else (RGBColor(0xf3, 0xf0, 0xfb) if plan['color'] == PURPLE else (RGBColor(0xff, 0xed, 0xe0) if plan['color'] == ORANGE else RGBColor(0xfe, 0xf9, 0xe7)))
    set_cell_bg(c, LIGHT_GREY)
    price_text = f"{plan['price_mo']} / mo\n({plan['price_yr']} / mo billed annually)"
    if plan['name'] == 'Enterprise':
        price_text = 'Custom pricing\nContact our team'
    cell_para(c, price_text, bold=False, size=9, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    c.width = Cm(3.8)

# Card features rows
max_features = max(len(p['highlights']) for p in plans_card_data)
feat_tbl = doc.add_table(rows=max_features + 1, cols=4)
feat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, plan in enumerate(plans_card_data):
    for j, feat in enumerate(plan['highlights']):
        c = feat_tbl.rows[j].cells[i]
        set_cell_bg(c, LIGHT_GREY if j % 2 == 0 else WHITE)
        cell_para(c, f'✓  {feat}', bold=False, size=8, color=DARK_GREY)
        c.width = Cm(3.8)
    # Fill empty rows
    for j in range(len(plan['highlights']), max_features):
        c = feat_tbl.rows[j].cells[i]
        set_cell_bg(c, LIGHT_GREY if j % 2 == 0 else WHITE)
        cell_para(c, '', size=8, color=WHITE)
        c.width = Cm(3.8)
    # CTA row
    c = feat_tbl.rows[max_features].cells[i]
    set_cell_bg(c, plan['color'])
    cell_para(c, plan['cta'], bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    c.width = Cm(3.8)

doc.add_paragraph()
add_para(doc, 'All plans include: 14-day free trial · TLS encryption · GDPR-ready · Hosted in Africa · 99.5%+ uptime',
         size=8, color=MID_GREY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2)

# Endpoints for pricing page
add_para(doc, 'Endpoints for this page:', bold=True, size=9, color=NAVY, space_before=10)
ep2 = doc.add_table(rows=4, cols=3)
pricing_eps = [
    ('GET',  '/api/v1/plans',                           'Load all public plans, pricing, and features'),
    ('GET',  '/api/v1/plans/compare',                   'Load full feature comparison matrix'),
    ('GET',  '/api/v1/checkout/active-sale',            'Check for active auto-apply sale campaign'),
    ('POST', '/api/v1/promotions/validate',             'Validate promo code before checkout'),
]
for i, (method, path, desc) in enumerate(pricing_eps):
    mc = ep2.rows[i].cells[0]
    pc = ep2.rows[i].cells[1]
    dc = ep2.rows[i].cells[2]
    m_color = GREEN if method == 'GET' else BLUE
    set_cell_bg(mc, m_color)
    cell_para(mc, method, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(pc, NAVY)
    cell_para(pc, path, bold=False, size=8, color=WHITE)
    set_cell_bg(dc, ROW_ALT if i%2==0 else WHITE)
    cell_para(dc, desc, bold=False, size=8, color=DARK_GREY)
    mc.width = Cm(1.5)
    pc.width = Cm(7.0)
    dc.width = Cm(8.0)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — ORG USER: PLAN COMPARISON TABLE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '5.  Org User — Plan Comparison Table', level=2, color=NAVY)
add_divider(doc)
add_para(doc, 'The full feature comparison table shown on the pricing page below the plan cards. Endpoint: GET /api/v1/plans/compare', size=9, color=MID_GREY, space_before=2, space_after=8)

comp_headers = ['Feature', 'Starter\n$15/mo', 'Professional\n$49/mo', 'Business\n$149/mo', 'Enterprise\nCustom']
comp_col_colors = [DARK_GREY, BLUE, PURPLE, ORANGE, GOLD]

comp_tbl = doc.add_table(rows=1, cols=5)
comp_tbl.style = 'Table Grid'
comp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(comp_headers):
    c = comp_tbl.rows[0].cells[j]
    set_cell_bg(c, comp_col_colors[j])
    cell_para(c, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

comp_rows = [
    ('LIMITS',                  '',    '',    '',    ''),
    ('Team members',           '5',   '25',  '100', '∞'),
    ('Projects',               '3',   '15',  '∞',   '∞'),
    ('Submissions / month',    '500', '5K',  '∞',   '∞'),
    ('SMS / month',            '200', '2K',  '10K', '∞'),
    ('API calls / month',      '2K',  '10K', '100K','∞'),
    ('Storage',                '5 GB','25 GB','100 GB','∞'),
    ('QR codes / month',       '50',  '500', '∞',   '∞'),
    ('Staff profiles',         '—',   '100', '∞',   '∞'),
    ('Uptime SLA',             '99.5%','99.9%','99.95%','99.99%'),
    ('CHANNELS',                '',    '',    '',    ''),
    ('SMS Feedback',           TICK,  TICK,  TICK,  TICK),
    ('WhatsApp Feedback',      DASH,  TICK,  TICK,  TICK),
    ('Phone Call AI (IVR)',    DASH,  DASH,  TICK,  TICK),
    ('Push Notifications',     TICK,  TICK,  TICK,  TICK),
    ('WhatsApp Notifications', DASH,  TICK,  TICK,  TICK),
    ('AI & INTELLIGENCE',       '',    '',    '',    ''),
    ('AI Conversation',        DASH,  TICK,  TICK,  TICK),
    ('Voice Transcription',    DASH,  TICK,  TICK,  TICK),
    ('AI Insights',            DASH,  DASH,  TICK,  TICK),
    ('ML Escalation Predictor',DASH,  DASH,  TICK,  TICK),
    ('Spark Streaming',        DASH,  DASH,  TICK,  TICK),
    ('Smart Recommendations',  TICK,  TICK,  TICK,  TICK),
    ('ANALYTICS & REPORTS',     '',    '',    '',    ''),
    ('Advanced Analytics',     DASH,  TICK,  TICK,  TICK),
    ('Custom Reports',         DASH,  TICK,  TICK,  TICK),
    ('INTEGRATIONS',            '',    '',    '',    ''),
    ('REST API Access',        TICK,  TICK,  TICK,  TICK),
    ('Webhooks',               DASH,  TICK,  TICK,  TICK),
    ('OAuth2 / Widget Embed',  TICK,  TICK,  TICK,  TICK),
    ('Audit Logs',             DASH,  TICK,  TICK,  TICK),
    ('SSO + SCIM',             DASH,  DASH,  DASH,  TICK),
    ('PLATFORM',                '',    '',    '',    ''),
    ('Dedicated Support',      DASH,  TICK,  TICK,  TICK),
    ('White-Label',            DASH,  DASH,  DASH,  TICK),
    ('Custom SLA contract',    DASH,  DASH,  DASH,  TICK),
    ('Price',                  '$15/mo','$49/mo','$149/mo','Custom'),
    ('CTA',                    'Start trial','Start trial','Start trial','Contact us'),
]

for i, (label, s, p, b, e) in enumerate(comp_rows):
    r = comp_tbl.add_row()
    is_section = s == '' and p == '' and b == '' and e == ''
    if is_section:
        for j in range(5):
            set_cell_bg(r.cells[j], DARK_GREY)
        cell_para(r.cells[0], label, bold=True, size=9, color=WHITE)
        for j in range(1, 5):
            cell_para(r.cells[j], '', size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        row_bg = LIGHT_GREY if i % 2 == 0 else WHITE
        set_cell_bg(r.cells[0], row_bg)
        cell_para(r.cells[0], label, bold=False, size=8.5, color=DARK_GREY)
        for j, val in enumerate([s, p, b, e], 1):
            is_tick = val == TICK
            is_price = label in ('Price', 'CTA')
            bg = GREEN_LIGHT if is_tick else row_bg
            set_cell_bg(r.cells[j], bg)
            cell_para(r.cells[j], val, bold=is_price or is_tick, size=8.5,
                      color=GREEN if is_tick else (MID_GREY if val == DASH else DARK_GREY),
                      align=WD_ALIGN_PARAGRAPH.CENTER)

for j, w in enumerate([5.0, 2.4, 2.4, 2.4, 2.4]):
    set_col_width(comp_tbl, j, w)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — ORG USER: CHECKOUT STEP-BY-STEP
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '6.  Org User — Checkout Step-by-Step', level=2, color=NAVY)
add_divider(doc)
add_para(doc, 'Complete journey from visiting the pricing page to an active subscription. Each step shows the screen, user action, API call, and expected outcome.', size=9, color=MID_GREY, space_before=2, space_after=10)

steps = [
    {
        'step': '1',
        'screen': 'Pricing Page',
        'action': 'User visits /pricing. Sees 4 plan cards (Starter, Professional, Business, Enterprise). Billing toggle (Monthly / Annual). Active sale banner auto-shown if applicable.',
        'endpoint': 'GET /api/v1/plans',
        'method': 'GET',
        'response': '{ "plans": [ {slug, display_name, pricing, limits, features} ] }',
        'outcome': 'Plans rendered. User selects billing cycle using toggle.',
    },
    {
        'step': '2',
        'screen': 'Pricing Page — Sale Banner',
        'action': 'App checks for an active auto-apply sale campaign and shows a banner if one is running.',
        'endpoint': 'GET /api/v1/checkout/active-sale',
        'method': 'GET',
        'response': '{ "active": true, "banner_text": "Launch Week — 40% OFF", "discount_value": 40 }',
        'outcome': 'Green banner appears above plan cards.',
    },
    {
        'step': '3',
        'screen': 'Pricing Page — Promo Code',
        'action': 'User enters a promo code in the "Have a promo code?" field. App validates before checkout.',
        'endpoint': 'POST /api/v1/promotions/validate',
        'method': 'POST',
        'response': '{ "valid": true, "discount_label": "30% off for 3 months" }',
        'outcome': 'Discount preview shown: "LAUNCH2026 — 30% off first 3 months".',
    },
    {
        'step': '4',
        'screen': 'Pricing Page — Plan Selection',
        'action': 'User clicks "Start free trial" on Professional card.',
        'endpoint': '— (client-side navigation)',
        'method': '—',
        'response': '—',
        'outcome': 'User is navigated to checkout/registration. Plan slug is passed as query param: /checkout?plan=professional&cycle=monthly',
    },
    {
        'step': '5',
        'screen': 'Login / Register',
        'action': 'If not logged in: user registers or logs in. Token with org_id is obtained.',
        'endpoint': 'POST /api/v1/auth/login\nPOST /api/v1/auth/login/verify-otp',
        'method': 'POST',
        'response': '{ "access_token": "eyJ...", "org_id": "uuid" }',
        'outcome': 'User is authenticated. org_id is present in JWT.',
    },
    {
        'step': '6',
        'screen': 'Subscription Check',
        'action': 'App checks if org already has a subscription (redirect if already subscribed).',
        'endpoint': 'GET /api/v1/subscriptions/current',
        'method': 'GET',
        'response': '{ "has_subscription": false } or { "status": "trialing" }',
        'outcome': 'If no subscription → show checkout. If trialing → show upgrade checkout. If active → redirect to dashboard.',
    },
    {
        'step': '7',
        'screen': 'Order Summary / Billing Preview',
        'action': 'App fetches exact price breakdown including active sale and promo code. Shown to user before payment.',
        'endpoint': 'POST /api/v1/subscriptions/billing-preview',
        'method': 'POST',
        'response': '{ "summary": { "subtotal": "$49.00", "discount": "-$14.70", "total": "$34.30" }, "line_items": [...] }',
        'outcome': 'User sees full breakdown: plan price, discount applied, total due today.',
    },
    {
        'step': '8',
        'screen': 'Payment Method Selection',
        'action': 'User selects payment provider: M-Pesa, AzamPay, Selcom, Airtel, Yas (mobile money) or PayPal (card/international) or Bank Transfer.',
        'endpoint': '— (client-side UI)',
        'method': '—',
        'response': '—',
        'outcome': 'Provider is selected. For mobile money: phone number input shown.',
    },
    {
        'step': '9',
        'screen': 'Checkout',
        'action': 'User clicks "Pay now" / "Start trial". App sends checkout request.',
        'endpoint': 'POST /api/v1/checkout',
        'method': 'POST',
        'response': '{ "subscription_id": "...", "payment_id": "...", "checkout_url": null (mpesa) or "https://paypal.com/..." }',
        'outcome': 'Payment initiated. Response varies by provider (see Section 7).',
    },
    {
        'step': '10',
        'screen': 'Payment Pending',
        'action': 'App polls for payment confirmation every 3 seconds (up to 60 seconds).',
        'endpoint': 'GET /api/v1/checkout/status/{payment_id}',
        'method': 'GET',
        'response': '{ "paid": false } → … → { "paid": true, "subscription_active": true }',
        'outcome': 'When paid=true: subscription is active, org is payment-verified.',
    },
    {
        'step': '11',
        'screen': 'Success Screen',
        'action': 'App refreshes subscription state and feature entitlements. Success screen shown.',
        'endpoint': 'GET /api/v1/subscriptions/my/features\nGET /api/v1/orgs/{slug}/badge',
        'method': 'GET',
        'response': '{ features: [...], limits: [...] }  +  { badge: { label: "Active Subscriber", color: "green" } }',
        'outcome': 'All plan features unlocked. Green "Active Subscriber" badge shown. User navigated to dashboard.',
    },
]

step_colors = [BLUE, GREEN, PURPLE, ORANGE, NAVY]

for s in steps:
    step_tbl = doc.add_table(rows=1, cols=2)
    step_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    sc = step_tbl.rows[0].cells[0]
    detail_c = step_tbl.rows[0].cells[1]
    sc.width = Cm(1.5)
    detail_c.width = Cm(15.0)

    color_idx = (int(s['step']) - 1) % len(step_colors)
    set_cell_bg(sc, step_colors[color_idx])
    cell_para(sc, f"Step\n{s['step']}", bold=True, size=12, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_bg(detail_c, LIGHT_GREY if int(s['step']) % 2 == 0 else WHITE)
    p = detail_c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{s['screen']}")
    run.bold = True; run.font.size = Pt(10); run.font.color.rgb = NAVY

    def add_cell_line(cell, label, value, label_color=MID_GREY, val_color=DARK_GREY, val_bold=False):
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        r1 = p2.add_run(f'{label}: ')
        r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = label_color
        r2 = p2.add_run(value)
        r2.bold = val_bold; r2.font.size = Pt(8.5); r2.font.color.rgb = val_color

    add_cell_line(detail_c, 'Screen', s['screen'], val_color=NAVY, val_bold=True)
    add_cell_line(detail_c, 'User action', s['action'])
    m_color = GREEN if s['method'] == 'GET' else (BLUE if s['method'] == 'POST' else MID_GREY)
    add_cell_line(detail_c, 'API call', f"[{s['method']}]  {s['endpoint']}", val_color=NAVY if s['method'] != '—' else MID_GREY, val_bold=(s['method'] != '—'))
    add_cell_line(detail_c, 'Response', s['response'], val_color=MID_GREY)
    add_cell_line(detail_c, 'Outcome', s['outcome'], val_color=GREEN if 'active' in s['outcome'].lower() or 'unlocked' in s['outcome'].lower() or 'success' in s['outcome'].lower() else DARK_GREY, val_bold=True if 'active' in s['outcome'].lower() else False)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — PAYMENT PROVIDER FLOWS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '7.  Org User — Payment Provider Flows', level=2, color=NAVY)
add_divider(doc)

providers = [
    {
        'name': 'M-Pesa / AzamPay / Selcom / Airtel / Yas',
        'color': GREEN,
        'subtitle': 'Mobile Money (Tanzania)',
        'request': '{ "plan_id": "uuid", "billing_cycle": "monthly", "provider": "mpesa",\n  "phone_number": "+255712345678", "payer_name": "John Komba", "promo_code": "LAUNCH2026" }',
        'response': '{ "payment_id": "pay-uuid", "payment": { "status": "pending",\n  "message": "USSD prompt sent. Enter your PIN." } }',
        'flow': [
            'POST /api/v1/checkout  →  response has payment_id',
            'Show screen: "Check your phone for M-Pesa prompt"',
            'Poll GET /api/v1/checkout/status/{payment_id} every 3s',
            'User receives USSD prompt, enters PIN on their phone',
            'Poll returns paid=true  →  navigate to success screen',
            'Timeout after 20 × 3s = 60s  →  show "Try again" screen',
        ],
    },
    {
        'name': 'PayPal',
        'color': BLUE,
        'subtitle': 'International (USD) — Visa / Mastercard / PayPal Balance',
        'request': '{ "plan_id": "uuid", "billing_cycle": "annual", "provider": "paypal",\n  "payer_email": "john@hospital.co.tz", "payer_name": "MNH" }',
        'response': '{ "payment_id": "pay-uuid", "checkout_url": "https://paypal.com/checkoutnow?token=...",\n  "payment": { "status": "pending" } }',
        'flow': [
            'POST /api/v1/checkout  →  response has checkout_url',
            'Open checkout_url in browser tab or WebView',
            'User logs into PayPal and approves payment',
            'PayPal redirects to your return_url callback',
            'Poll GET /api/v1/checkout/status/{payment_id}',
            'paid=true  →  close WebView  →  show success screen',
        ],
    },
    {
        'name': 'Bank Transfer',
        'color': GOLD,
        'subtitle': 'Wire / CRDB / NMB — Manual Confirmation',
        'request': '{ "plan_id": "uuid", "billing_cycle": "monthly", "provider": "bank_transfer" }',
        'response': '{ "invoice": { "invoice_number": "INV-2026-001234", "total_usd": "149.00" },\n  "payment": { "instructions": "Transfer USD 149.00. Reference: INV-2026-001234." } }',
        'flow': [
            'POST /api/v1/checkout  →  response has invoice + instructions',
            'Show bank details: account, SWIFT, reference number',
            'User makes the bank transfer (1–2 business days)',
            'Admin confirms via PATCH /api/v1/admin/orgs/{id}/payment-verification',
            'Subscription activates  →  user gets email confirmation',
            'GET /api/v1/subscriptions/current  →  status = active',
        ],
    },
]

for prov in providers:
    add_heading(doc, prov['name'], level=3, color=prov['color'], space_before=12, space_after=4)
    add_para(doc, prov['subtitle'], size=9, color=MID_GREY, space_before=0, space_after=6)

    prov_tbl = doc.add_table(rows=2, cols=2)
    prov_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Request cell
    rc = prov_tbl.rows[0].cells[0]
    set_cell_bg(rc, NAVY)
    cell_para(rc, 'REQUEST BODY', bold=True, size=8, color=WHITE)
    p_req = rc.add_paragraph()
    p_req.paragraph_format.space_before = Pt(2)
    rr = p_req.add_run(prov['request'])
    rr.font.size = Pt(7.5); rr.font.color.rgb = RGBColor(0xb0, 0xc4, 0xde)
    rr.font.name = 'Courier New'

    # Response cell
    resp_c = prov_tbl.rows[0].cells[1]
    set_cell_bg(resp_c, DARK_GREY)
    cell_para(resp_c, 'RESPONSE', bold=True, size=8, color=WHITE)
    p_resp = resp_c.add_paragraph()
    rresp = p_resp.add_run(prov['response'])
    rresp.font.size = Pt(7.5); rresp.font.color.rgb = GREEN_LIGHT
    rresp.font.name = 'Courier New'

    prov_tbl.rows[0].cells[0].width = Cm(8.0)
    prov_tbl.rows[0].cells[1].width = Cm(8.5)

    # Flow steps cell (spans both cols via merge)
    flow_c = prov_tbl.rows[1].cells[0].merge(prov_tbl.rows[1].cells[1])
    set_cell_bg(flow_c, LIGHT_GREY)
    cell_para(flow_c, 'FRONTEND FLOW', bold=True, size=8, color=prov['color'])
    for fi, step_text in enumerate(prov['flow'], 1):
        fp = flow_c.add_paragraph()
        fp.paragraph_format.space_before = Pt(1)
        fp.paragraph_format.space_after = Pt(1)
        fr = fp.add_run(f'  {fi}.  {step_text}')
        fr.font.size = Pt(8.5)
        fr.font.color.rgb = DARK_GREY
        fr.bold = ('paid=true' in step_text or 'success' in step_text.lower() or 'active' in step_text)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — POST-PAYMENT & VERIFICATION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '8.  Org User — Post-Payment & Verification', level=2, color=NAVY)
add_divider(doc)
add_para(doc, 'What happens automatically when paid=true, and what the org sees on their dashboard.', size=9, color=MID_GREY, space_before=2, space_after=8)

post_tbl = doc.add_table(rows=1, cols=4)
post_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Event', 'Set by', 'How frontend sees it', 'Badge']):
    c = post_tbl.rows[0].cells[j]
    set_cell_bg(c, NAVY)
    cell_para(c, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

post_rows = [
    ('Subscription status = active',    'subscription_service (auto)',    'GET /subscriptions/current → status: "active"',           '—'),
    ('is_payment_verified = true',      'auth_service (auto)',             'GET /orgs/{slug}/badge → is_payment_verified: true',      'Green "Active Subscriber"'),
    ('All plan features unlocked',      'subscription_service (auto)',     'GET /subscriptions/my/features → all plan flags enabled', '—'),
    ('Feature gate map rebuilt',        'Frontend (invalidate cache)',     'Re-fetch /subscriptions/my/features',                    '—'),
    ('KYC submitted (optional, later)', 'Org submits docs',               'POST /orgs/my/kyc/submit',                               '—'),
    ('KYC approved (optional, later)',  'Platform admin approves',        'GET /orgs/{slug}/badge → is_kyc_verified: true',          'Blue "Verified Business"'),
]

for i, (evt, set_by, how, badge) in enumerate(post_rows):
    r = post_tbl.add_row()
    row_bg = GREEN_LIGHT if i < 3 else (BLUE_LIGHT if i == 5 else (LIGHT_GREY if i % 2 == 0 else WHITE))
    for j in range(4):
        set_cell_bg(r.cells[j], row_bg)
    cell_para(r.cells[0], evt, bold=True, size=8.5, color=DARK_GREY)
    cell_para(r.cells[1], set_by, bold=False, size=8, color=MID_GREY)
    cell_para(r.cells[2], how, bold=False, size=8, color=NAVY)
    cell_para(r.cells[3], badge, bold=(badge != '—'), size=8.5,
              color=GREEN if 'Green' in badge else (BLUE if 'Blue' in badge else MID_GREY))
for j, w in enumerate([4.5, 3.5, 5.5, 3.0]):
    set_col_width(post_tbl, j, w)

doc.add_paragraph()

# Badge design mockup
add_heading(doc, 'Verification Badge States', level=3, color=NAVY, space_before=12, space_after=6)

badge_tbl = doc.add_table(rows=4, cols=3)
badge_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['State', 'Badge Appearance', 'When shown']):
    c = badge_tbl.rows[0].cells[j]
    set_cell_bg(c, DARK_GREY)
    cell_para(c, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

badge_rows = [
    ('No badge',              '(nothing rendered)',                         'is_payment_verified=false AND is_kyc_verified=false'),
    ('Active Subscriber',     '✓  Active Subscriber',     GREEN,           'is_payment_verified=true, is_kyc_verified=false'),
    ('Verified Business',     '✓  Verified Business',     BLUE,            'is_kyc_verified=true (takes priority over green)'),
]

for i, row_d in enumerate(badge_rows):
    r = badge_tbl.rows[i+1]
    set_cell_bg(r.cells[0], LIGHT_GREY if i%2==0 else WHITE)
    cell_para(r.cells[0], row_d[0], bold=True, size=9, color=DARK_GREY)

    if len(row_d) == 4:
        set_cell_bg(r.cells[1], row_d[2])
        cell_para(r.cells[1], row_d[1], bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(r.cells[2], LIGHT_GREY if i%2==0 else WHITE)
        cell_para(r.cells[2], row_d[3], bold=False, size=8.5, color=DARK_GREY)
    else:
        set_cell_bg(r.cells[1], LIGHT_GREY)
        cell_para(r.cells[1], row_d[1], bold=False, size=9, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        set_cell_bg(r.cells[2], LIGHT_GREY if i%2==0 else WHITE)
        cell_para(r.cells[2], row_d[2], bold=False, size=8.5, color=DARK_GREY)

for j, w in enumerate([3.5, 4.5, 8.5]):
    set_col_width(badge_tbl, j, w)

add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — ENDPOINT REFERENCE PER SCREEN
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '9.  Endpoint Reference per Screen', level=2, color=NAVY)
add_divider(doc)
add_para(doc, 'Every API call mapped to the screen that triggers it. All paths prefixed with /api/v1. Service shown for internal routing.', size=9, color=MID_GREY, space_before=2, space_after=10)

screen_groups = [
    ('Pricing Page',  'Public — no auth required', [
        ('GET',   '/plans',                        'Load all public plans with pricing, limits, features',  'subscription_service'),
        ('GET',   '/plans/compare',                'Feature comparison matrix (for comparison table)',      'subscription_service'),
        ('GET',   '/plans/features',               'Full feature catalog (46 features with descriptions)', 'subscription_service'),
        ('GET',   '/plans/addons',                 'Available add-ons and pricing',                        'subscription_service'),
        ('GET',   '/checkout/active-sale',         'Active auto-apply sale campaign (banner)',              'subscription_service'),
        ('POST',  '/promotions/validate',          'Validate promo code, show discount preview',           'subscription_service'),
    ]),
    ('Authentication (Login / Register)', 'Required before checkout', [
        ('POST',  '/auth/register',                'Create account (new user)',                             'auth_service'),
        ('POST',  '/auth/login',                   'Log in — returns login_token',                         'auth_service'),
        ('POST',  '/auth/login/verify-otp',        'Submit OTP — returns access_token + org_id',          'auth_service'),
        ('POST',  '/auth/switch-org',              'Switch org context (multi-org users)',                 'auth_service'),
    ]),
    ('Checkout / Order Summary', 'Auth required — org_id in JWT', [
        ('GET',   '/subscriptions/current',        'Check existing subscription before checkout',          'subscription_service'),
        ('POST',  '/subscriptions/billing-preview','Order summary: exact price, tax, discounts',           'subscription_service'),
        ('POST',  '/checkout',                     'Initiate payment (all providers)',                     'subscription_service'),
        ('GET',   '/checkout/status/{payment_id}', 'Poll payment status (every 3s)',                       'subscription_service'),
        ('POST',  '/checkout/pay-invoice/{id}',    'Pay an overdue invoice',                               'subscription_service'),
    ]),
    ('Dashboard — After Login', 'Auth required', [
        ('GET',   '/subscriptions/current',        'Load subscription state on app start',                 'subscription_service'),
        ('GET',   '/subscriptions/my/features',    'Load feature gate map (what org can access)',          'subscription_service'),
        ('GET',   '/orgs/my/verification',         'Check both verification tracks',                      'auth_service'),
        ('GET',   '/orgs/{slug}/badge',            'Public badge for display (no auth needed)',            'auth_service'),
    ]),
    ('Account Settings — Billing', 'Auth required', [
        ('GET',   '/subscriptions/current',        'Current plan, status, renewal date',                   'subscription_service'),
        ('GET',   '/subscriptions/invoices',       'Billing history',                                      'subscription_service'),
        ('POST',  '/subscriptions/upgrade',        'Upgrade to higher plan (prorated)',                    'subscription_service'),
        ('POST',  '/subscriptions/downgrade',      'Downgrade (effective next billing cycle)',             'subscription_service'),
        ('POST',  '/subscriptions/switch-billing-cycle', 'Monthly ↔ Annual',                              'subscription_service'),
        ('POST',  '/subscriptions/apply-promo',   'Apply promo code to active subscription',              'subscription_service'),
        ('POST',  '/subscriptions/pause',          'Pause subscription (Business/Enterprise)',             'subscription_service'),
        ('POST',  '/subscriptions/resume',         'Resume paused subscription',                          'subscription_service'),
        ('POST',  '/subscriptions/cancel',         'Cancel at period end or immediately',                  'subscription_service'),
        ('GET',   '/subscriptions/events',         'Subscription audit timeline',                         'subscription_service'),
    ]),
    ('Account Settings — Verification / KYC', 'Auth required', [
        ('GET',   '/orgs/my/verification',         'Both verification tracks + latest KYC summary',       'auth_service'),
        ('POST',  '/orgs/my/kyc/submit',           'Create KYC submission',                               'auth_service'),
        ('POST',  '/orgs/my/kyc/documents/upload', 'Upload document file (multipart)',                    'auth_service'),
        ('POST',  '/orgs/my/kyc/documents',        'Add document by URL to existing submission',          'auth_service'),
        ('DELETE','/orgs/my/kyc/documents/{id}',   'Remove document (pending/more_info only)',            'auth_service'),
        ('GET',   '/orgs/my/kyc',                  'KYC history with documents (last 5 submissions)',     'auth_service'),
    ]),
    ('Public Org / Product Pages', 'No auth required', [
        ('GET',   '/orgs/{slug}/badge',            'Verification badge (color, label)',                    'auth_service'),
        ('GET',   '/plans/{slug}',                 'Single plan detail by slug',                           'subscription_service'),
    ]),
]

method_colors = {'GET': GREEN, 'POST': BLUE, 'DELETE': RGBColor(0xdc, 0x26, 0x26), 'PATCH': PURPLE}

for screen_name, auth_note, endpoints in screen_groups:
    # Screen group header
    sg = doc.add_table(rows=1, cols=1)
    sg_c = sg.rows[0].cells[0]
    set_cell_bg(sg_c, NAVY)
    cell_para(sg_c, f'  {screen_name}  ·  {auth_note}', bold=True, size=10, color=WHITE)

    ep_tbl = doc.add_table(rows=len(endpoints), cols=4)
    ep_tbl.style = 'Table Grid'
    for i, (method, path, desc, service) in enumerate(endpoints):
        row_bg = LIGHT_GREY if i % 2 == 0 else WHITE
        mc = ep_tbl.rows[i].cells[0]
        pc = ep_tbl.rows[i].cells[1]
        dc = ep_tbl.rows[i].cells[2]
        sc = ep_tbl.rows[i].cells[3]
        set_cell_bg(mc, method_colors.get(method, DARK_GREY))
        cell_para(mc, method, bold=True, size=8, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(pc, DARK_GREY if i%2==0 else RGBColor(0x4b, 0x55, 0x63))
        cell_para(pc, path, bold=False, size=8, color=WHITE)
        set_cell_bg(dc, row_bg)
        cell_para(dc, desc, bold=False, size=8, color=DARK_GREY)
        set_cell_bg(sc, row_bg)
        cell_para(sc, service, bold=False, size=7.5, color=MID_GREY, italic=True)
        mc.width = Cm(1.5)
        pc.width = Cm(6.5)
        dc.width = Cm(7.0)
        sc.width = Cm(2.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ── Back cover ────────────────────────────────────────────────────────────────
add_page_break(doc)
add_heading(doc, 'Riviwa Platform', level=1, size=22, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=80, space_after=6)
add_para(doc, 'Grievance & Feedback Management', size=12, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_before=0)
doc.add_paragraph()
add_para(doc, 'api.riviwa.com  ·  billing@riviwa.com', size=10, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)
add_para(doc, 'Subscription Implementation Guide  ·  Version 2.5  ·  May 2026', size=9, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──────────────────────────────────────────────────────────────────────
output = r"C:\Users\mmadi\OneDrive\Documents\FinalRiviwa\Riviwa\Riviwa_Subscription_Pricing_Design.docx"
doc.save(output)
print(f"Saved: {output}")
